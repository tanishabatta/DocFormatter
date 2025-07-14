import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx.oxml.ns import qn



import shutil
import getpass
import os




input_folder = "input_docs"
output_folder = "formatted_docs"
image_path = "Picture1.png"  

# Ensure output folder exists
os.makedirs(output_folder, exist_ok=True)

#Main title

#Creates a title that is correctly formatted for the document

def main_title(filename):
    
    base_name = os.path.splitext(filename)[0]
    words = base_name.split('_')
    formatted_words = [
        word if word[0].isupper() 
        else word.capitalize() 
        for word in words
    ]
    return ' '.join(formatted_words)

#remove borders
def remove_table_borders(table):
   
    tbl = table._tbl
    tblPr = tbl.tblPr

    # Remove any existing border settings
    for border_tag in ["tblBorders", "tblBorder"]:
        el = tblPr.find(qn(f"w:{border_tag}"))
        if el is not None:
            tblPr.remove(el)

#Title and Image table

#Adds the main title and the sagenet image
def title_table(doc, title_text, image_path):
    table = doc.add_table(rows=1, cols=2)
   
    remove_table_borders(table)
   
    table.auto_fit = False
    table.columns[0].width = Inches(6)
    table.columns[1].width = Inches(6)

    # Left cell: title
    cell1 = table.cell(0, 0) #column 1 row 1
    p1 = cell1.paragraphs[0] 
    run1 = p1.add_run(title_text)
    run1.bold = True
    run1.font.size = Pt(20)
    run1.font.color.rgb = RGBColor(0, 120, 173)  
    run1.font.name = 'Calibri'

    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Right cell: image
    cell2 = table.cell(0, 1)
    p2 = cell2.paragraphs[0]
    try:
        run2 = p2.add_run()
        run2.add_picture(image_path, width=Inches(1.5))
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    except Exception:
        p2.add_run(f"[Image missing: {image_path}]")
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Move table to top of document
    doc._body._element.insert(0, table._element)

"""
def title_with_top_right_image(doc, title_text, image_path):
    
    Adds a title aligned to the left and an image aligned to the top-right without using a table.
    
    # First paragraph: Image aligned to right
    image_paragraph = doc.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_img = image_paragraph.add_run()
    try:
        run_img.add_picture(image_path, width=Inches(1.5))
    except Exception:
        image_paragraph.add_run(f"[Image not found: {image_path}]")

    # Second paragraph: Title aligned to left
    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_title = title_paragraph.add_run(title_text)
    run_title.bold = True
    run_title.font.size = Pt(20)
    run_title.font.color.rgb = RGBColor(0, 120, 173)
    run_title.font.name = 'Calibri'

    # Insert image and title paragraphs at the top
    doc._body._element.insert(0, title_paragraph._element)
    doc._body._element.insert(0, image_paragraph._element)


"""

#Purpose and Author

def purpose_author(doc, purpose_text, author_text):
    # Create a 2-row, 1-column table
    table = doc.add_table(rows=2, cols=1)
    remove_table_borders(table)
    
    table.allow_autofit = True
    

    # Row 1 — Purpose line
    cell1 = table.cell(0, 0)
    para1 = cell1.paragraphs[0]
    para1.clear()

    run_purpose_label = para1.add_run("Purpose: ")
    run_purpose_label.bold = True
    run_purpose_label.underline = True
    run_purpose_label.font.size = Pt(11)

    run_purpose_text = para1.add_run(purpose_text)
    run_purpose_text.font.size = Pt(11)

    # Row 2 — Author line
    cell2 = table.cell(1, 0)
    para2 = cell2.paragraphs[0]
    para2.clear()

    run_author_label = para2.add_run("Author: ")
    run_author_label.bold = True
    run_author_label.underline = True
    run_author_label.font.size = Pt(11)

    run_author_text = para2.add_run(author_text)
    run_author_text.font.size = Pt(11)

    # === Insert below the first table ===
    body = doc._body._element
    tbls = body.findall('.//w:tbl', body.nsmap)

    if len(tbls) >= 1:
        body.remove(table._element)  # It's added at end by default
        body.insert(body.index(tbls[0]) + 1, table._element)

    return table
    



    # Move table to top of document
    doc._body._element.insert(0, table._element)




    



#Border
def add_page_borders(doc):
        for section in doc.sections:
            sectPr = section._sectPr

        # Create the pgBorders element
            pgBorders = OxmlElement('w:pgBorders')
            pgBorders.set(qn('w:offsetFrom'), 'page')  # or 'text'

        # Border settings
            for side in ['top', 'left', 'bottom', 'right']:
                element = OxmlElement(f'w:{side}')
                element.set(qn('w:val'), 'single')      # Border style
                element.set(qn('w:sz'), '12')           # Width (1/8 pts)
                element.set(qn('w:space'), '24')        # Space between border and text
                element.set(qn('w:color'), 'D3D3D3')    # Border color (black)
                pgBorders.append(element)

        # Attach to section
            sectPr.append(pgBorders)

#Header
def add_header(doc):
    for section in doc.sections:
        header = section.header
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_parapgrahph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        paragraph.clear()

        run = paragraph.add_run("Page ")

        fld_char1 = OxmlElement('w:fldChar')
        fld_char1.set(qn('w:fldCharType'), 'begin')

        instr_text = OxmlElement('w:instrText')
        instr_text.set(qn('xml:space'), 'preserve')
        instr_text.text = "PAGE"

        fld_char2 = OxmlElement('w:fldChar')
        fld_char2.set(qn('w:fldCharType'), 'separate')

        fld_char3 = OxmlElement('w:t')
        fld_char3.text = "1"  # Placeholder, Word updates it dynamically

        fld_char4 = OxmlElement('w:fldChar')
        fld_char4.set(qn('w:fldCharType'), 'end')

        run._r.append(fld_char1)
        run._r.append(instr_text)
        run._r.append(fld_char2)
        run._r.append(fld_char3)
        run._r.append(fld_char4)

        # Add " of "
        paragraph.add_run(" of ")

        # Add NUMPAGES field
        run2 = paragraph.add_run()

        fld_char1_np = OxmlElement('w:fldChar')
        fld_char1_np.set(qn('w:fldCharType'), 'begin')

        instr_text_np = OxmlElement('w:instrText')
        instr_text_np.set(qn('xml:space'), 'preserve')
        instr_text_np.text = "NUMPAGES"

        fld_char2_np = OxmlElement('w:fldChar')
        fld_char2_np.set(qn('w:fldCharType'), 'separate')

        fld_char3_np = OxmlElement('w:t')
        fld_char3_np.text = "1"  # Placeholder

        fld_char4_np = OxmlElement('w:fldChar')
        fld_char4_np.set(qn('w:fldCharType'), 'end')

        run2._r.append(fld_char1_np)
        run2._r.append(instr_text_np)
        run2._r.append(fld_char2_np)
        run2._r.append(fld_char3_np)
        run2._r.append(fld_char4_np)

        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        run2.font.size = Pt(12)
        run2.font.name = 'Times New Roman'


#Footer


def add_footer(doc):
    
    today_str = datetime.today().strftime("%m/%d/%y")

    for section in doc.sections:
        footer = section.footer

        # Use the first paragraph or create one
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.clear()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Construct footer text
        footer_text = f"Revised By: Tanisha Batta      Rev: 1.1    Date: {today_str}"
        run = paragraph.add_run(footer_text)

        # Optional styling
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(128, 128, 128)  

#page break

def add_page_break_at_end(doc):
   
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)

#Release Notes title

def release_notes_title(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Release Notes")
    run.font.size = Pt(16)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(0, 120, 173)  # Sagenet blue color
    title.add_run("\n")  # Add a new line after the title

#Release Notes table

def set_cell_border(cell, border_color="808080", border_sz="4"):
   
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
   
    borders = parse_xml(rf'''
    <w:tcBorders {nsdecls('w')}>
      <w:top w:val="single" w:sz="{border_sz}" w:color="{border_color}" w:space="0"/>
      <w:left w:val="single" w:sz="{border_sz}" w:color="{border_color}" w:space="0"/>
      <w:bottom w:val="single" w:sz="{border_sz}" w:color="{border_color}" w:space="0"/>
      <w:right w:val="single" w:sz="{border_sz}" w:color="{border_color}" w:space="0"/>
    </w:tcBorders>
    ''')
   
    tcPr.append(borders)

def release_notes_table(doc):
    table = doc.add_table (rows=3, cols=4)

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Version'
    hdr_cells[1].text = 'Date'
    hdr_cells[2].text = 'Modified By'
    hdr_cells[3].text = 'Changes Made'

    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.rgb = RGBColor(0, 0, 0)
                run.font.name = 'Calibri'
                run.font.size = Pt(12)

    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, border_color="808080", border_sz="4")

    data = [
        ["1.0", "Unknown", "Unknown", "Initial Issue"],
        ["1.1", datetime.today().strftime("%m/%d/%y"), "Tanisha Batta", "Added to the template and adjusted formatting."]
    ]

    for row_idx, row_data in enumerate(data, start=1):
        for col_idx, value in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            cell.text = value
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    run.font.name = 'Calibri'

    # === Apply border to all cells ===
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, border_color="808080", border_sz="4")



def process_documents(input_folder, output_folder, image_path):
    for filename in os.listdir(input_folder):
        if filename.endswith(".docx"):
            input_path = os.path.join(input_folder, filename)
            base_name, ext = os.path.splitext(filename)
            new_filename = f"{base_name}_formatted{ext}"
            output_path = os.path.join(output_folder, new_filename)

            print(f"Processing: {filename}")

            
            doc = Document(input_path)

            
             # Create a clean title from filename
            title_text =main_title(filename)

            

            # Insert title + image table at the top
            title_table(doc, title_text, image_path=image_path)
            #title_with_image(doc, title_text, image_path=image_path)
            purpose_author(doc, purpose_text= "Unknown", author_text= "Unknown")
            
            add_page_borders(doc)
            add_header(doc)
            add_footer(doc)
            add_page_break_at_end(doc)
            release_notes_title(doc)
            release_notes_table(doc)
            



            
            doc.save(output_path)

    print("\n All documents formatted and saved to:", output_folder)


process_documents(input_folder, output_folder, image_path)
